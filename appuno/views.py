from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from .models import Registro
from .chatbot_logic import preguntas
from docx import Document

# Índice para rastrear la pregunta actual
pregunta_actual = 0

# Vista para renderizar el HTML del chatbot
def chatbot_view(request):
    return render(request, 'index.html')  # El archivo HTML debe estar en la carpeta templates

# API del chatbot que maneja las respuestas
def chatbot(request):
    global pregunta_actual
    
    if request.method == "POST":
        user_message = request.POST.get("message", "")

        # Obtener la pregunta actual
        pregunta = preguntas[pregunta_actual]  # Usamos el índice para obtener la pregunta correcta

        # Guardar la respuesta en la base de datos
        if user_message:  # Si hay una respuesta del usuario
            Registro.objects.create(pregunta=pregunta, respuesta=user_message)

            # Incrementar el índice para la siguiente pregunta
            pregunta_actual += 1

            # Si ya no hay más preguntas, reiniciamos el índice
            if pregunta_actual >= len(preguntas):
                pregunta_actual = 0  # O puedes hacer que el chatbot termine o reinicie de otra forma

        # Si no hay respuesta, preguntar de nuevo
        if pregunta_actual < len(preguntas):
            respuesta = preguntas[pregunta_actual]  # Enviar la siguiente pregunta
        else:
            respuesta = "¡Gracias por responder todas las preguntas!"

        return JsonResponse({"message": respuesta})

    return JsonResponse({"error": "Método no permitido"}, status=405)

def resumen(req): 
    return render(req, 'resumen.html')
def generar_resumen(request):
    # Crear el documento de Word
    doc = Document()
    doc.add_heading('Resumen del Análisis de Aceite', 0)

    # Obtener todas las respuestas del usuario desde la base de datos
    respuestas = Registro.objects.all()

    # Crear una tabla en el documento
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Encabezados de la tabla
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pregunta'
    hdr_cells[1].text = 'Respuesta'

    # Llenar la tabla con las respuestas
    for registro in respuestas:
        row_cells = table.add_row().cells
        row_cells[0].text = registro.pregunta
        row_cells[1].text = registro.respuesta

    # Crear la respuesta para la descarga del archivo
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=resumen_analisis_aceite.docx'
    
    # Guardar el documento en el objeto HttpResponse para enviarlo al usuario
    doc.save(response)

    return response
def reiniciar_chatbot(request):
    global pregunta_actual
    pregunta_actual = 0  # Reiniciamos el índice para volver a las preguntas iniciales
    return JsonResponse({"message": preguntas[pregunta_actual]})  # Enviamos la primera pregunta nuevamente